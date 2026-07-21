from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.rag import embeddings, ingest, vectorstore
from app.rag.loaders import Document


@pytest.fixture
def client():
    return vectorstore.get_client(location=":memory:")


@pytest.fixture(autouse=True)
def fake_embeddings(monkeypatch):
    """Deterministic 4-dim embeddings so tests need no API/network."""

    async def fake_embed(texts):
        return [[float(len(t)), 1.0, 0.0, -1.0] for t in texts]

    monkeypatch.setattr(embeddings, "embed", fake_embed)


async def test_ingest_document_stores_chunks(client):
    long_text = "\n\n".join(f"Paragraph {i} with content. " * 10 for i in range(30))
    doc = Document(text=long_text, metadata={"source": "notes.pdf", "filetype": "pdf"})

    result = await ingest.ingest_document(doc, client, collection="documents")

    assert result.num_chunks > 1
    assert len(result.ids) == result.num_chunks
    assert result.source == "notes.pdf"
    assert client.collection_exists("documents")


async def test_ingested_chunks_are_searchable(client):
    doc = Document(text="alpha beta gamma delta. " * 40, metadata={"source": "a.pdf"})
    await ingest.ingest_document(doc, client, collection="documents")

    hits = vectorstore.search(client, "documents", query_vector=[10.0, 1.0, 0.0, -1.0])
    assert hits
    assert "text" in hits[0].payload
    assert hits[0].payload["source"] == "a.pdf"


async def test_empty_document_stores_nothing(client):
    doc = Document(text="   ", metadata={"source": "empty.pdf"})
    result = await ingest.ingest_document(doc, client)
    assert result.num_chunks == 0
    assert result.ids == []


async def test_ingest_file_from_disk(client, tmp_path: Path):
    d = DocxDocument()
    for i in range(30):
        d.add_paragraph(f"Sentence number {i} about revenue and strategy.")
    p = tmp_path / "report.docx"
    d.save(p)

    result = await ingest.ingest_file(p, client, collection="documents")
    assert result.num_chunks >= 1
    assert result.source == "report.docx"
