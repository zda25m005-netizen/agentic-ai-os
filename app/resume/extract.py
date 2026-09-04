"""Extract plain text from an uploaded resume (PDF or DOCX). No network."""

from __future__ import annotations

import io


class UnsupportedResume(ValueError):
    """Raised when the file type is not a supported resume format."""


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(t for t in parts if t).strip()


def extract_resume_text(filename: str, data: bytes) -> str:
    """Dispatch on extension. Returns extracted text (may be empty)."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".docx"):
        return extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()
    raise UnsupportedResume("Only PDF, DOCX, or TXT resumes are supported.")
